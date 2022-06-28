import tkinter
import time

from tkinter import *
from tkinter import filedialog
from PIL import Image, ImageTk
import docx as dc
main= Tk ()
main.configure(bg='#f9fcfb')
main.title("SS MERGER")
#main.call('wm', 'iconphoto', main._w, PhotoImage(file='C:/Users/Dell/Downloads/myicon.ico'))
main.iconbitmap(r'icon.ico')

canvas= tkinter.Canvas(main,width=900, height=400)
canvas.grid(columnspan=6, rowspan=6)
#Add Logos
logo = Image.open('logo1.png')
logo = ImageTk.PhotoImage(logo)
logo_label = Label(image=logo, height=400, width=500)
logo_label.image = logo
logo_label.grid(row=0, column=3)
#Function for Merging SS
def SS_Merge():
    doc = dc.Document('C:/Users/Dell/Documents/Merge SS/SS merged.docx')
    doc1 = dc.Document(file1_name)
    doc2 = dc.Document(file2_name)
    doc3 = dc.Document(file3_name)
    List_ss1_col1 = []
    List_ss1_col2 = []
    List_ss2_col1 = []
    List_ss2_col2 = []
    List_ss3_col1 = []
    List_ss3_col2 = []
    # Lists to keep each files two column's values.

    for table1 in doc1.tables:
        for index in range(2, 20, 1):
            if (index == 8 or index == 13):
                List_ss1_col1.append('')
                List_ss1_col2.append('')
                continue

            # reading values from each cells for each of the six rows and appending to two column lists
            cell_id = table1.rows[index].cells
            List_ss1_col1.append(cell_id[1].text)
            List_ss1_col2.append(cell_id[2].text)
    print(List_ss1_col1)
    print(List_ss1_col2)

    for table2 in doc2.tables:
        for index in range(2, 20, 1):
            if (index == 8 or index == 13):
                List_ss2_col1.append('')
                List_ss2_col2.append('')
                continue
            # reading values for separate cells of second file and appending two column's data to two separate lists.
            cell_id = table2.rows[index].cells
            List_ss2_col1.append(cell_id[1].text)
            List_ss2_col2.append(cell_id[2].text)
    print(List_ss2_col1)
    print(List_ss2_col2)

    for table3 in doc3.tables:
        for index in range(2, 20, 1):
            if (index == 8 or index == 13):
                List_ss3_col1.append('')
                List_ss3_col2.append('')
                continue
            # reading and appending data to the list for 3rd file.
            cell_id = table3.rows[index].cells
            List_ss3_col1.append(cell_id[1].text)
            List_ss3_col2.append(cell_id[2].text)
    print(List_ss3_col1)
    print(List_ss3_col2)

    for table in doc.tables:
        for i in range(6):
            # Reading data from respective lists and writing them on specified locations in the merged file.
            column = table.rows[2 + i * 3].cells
            # data appending for the first row (row 2 for i=0)
            # data writing to the respective cells for each of the six position of the column list of each file
            column[1].text = List_ss1_col1[i]
            # data writing to the first column from first position of the list
            column[2].text = List_ss1_col2[i]
            # incrementing rows index by three each after a iteration of putting values from the each locations of list.
            column = table.rows[3 + i * 3].cells
            column[1].text = List_ss2_col1[i]
            column[2].text = List_ss2_col2[i]
            column = table.rows[4 + i * 3].cells
            column[1].text = List_ss3_col1[i]
            column[2].text = List_ss3_col2[i]
        for i in range(4):
            column = table.rows[21 + i * 3].cells
            column[1].text = List_ss1_col1[i + 7]
            column[2].text = List_ss1_col2[i + 7]
            column = table.rows[22 + i * 3].cells
            column[1].text = List_ss2_col1[i + 7]
            column[2].text = List_ss2_col2[i + 7]
            column = table.rows[23 + i * 3].cells
            column[1].text = List_ss3_col1[i + 7]
            column[2].text = List_ss3_col2[i + 7]
        for i in range(6):
            column = table.rows[34 + i * 3].cells
            column[1].text = List_ss1_col1[i + 12]
            column[2].text = List_ss1_col2[i + 12]
            column = table.rows[35 + i * 3].cells
            column[1].text = List_ss2_col1[i + 12]
            column[2].text = List_ss2_col2[i + 12]
            column = table.rows[36 + i * 3].cells
            column[1].text = List_ss3_col1[i + 12]
            column[2].text = List_ss3_col2[i + 12]
    doc.save('ss merged.docx')


def Merge():
    Instructions.config(text = "Merger Sucessful! MERGED FILE LOCATION:C:/Users/Dell/Documents/Merge SS/SS merged.docx")
    try:
        SS_Merge()
    except:
        Instructions.config(text="Error Merging Your Files. Please try with Other File Formats.")
    Instructions.after(6000, lambda: Instructions.config(text = "Select three SS files and Click Merge"))
    first_text.delete(1.0, 'end')
    second_text.delete(1.0, 'end')
    third_text.delete(1.0, 'end')

def First_browse():
    global file1_name
    first_text.delete(1.0,'end')
    first_text.config(state="normal")
    main.file1=filedialog.askopenfilename(parent=main, title="Choose a file", filetype=[("docx files","*.docx")])
    file1_name=main.file1
    first_text.insert(INSERT,main.file1)
def Second_browse():
    global file2_name
    second_text.delete(1.0, 'end')
    second_text.config(state="normal")
    main.file2=filedialog.askopenfilename(parent=main, title="Choose a file", filetype=[("docx files","*.docx")])
    second_text.insert(INSERT,main.file2)
    file2_name= main.file2
def third_browse():
    global file3_name
    third_text.delete(1.0, 'end')
    third_text.config(state="normal")
    main.file3=filedialog.askopenfilename(parent=main, title="Choose a file", filetype=[("docx files","*.docx")])
    third_text.insert(INSERT,main.file3)
    file3_name= main.file3
#Instructions
frame= LabelFrame(main)
frame.grid(column=3, row=1,padx=5,pady=5,rowspan=5)
#frame.place(relx=0.5,rely=0.2,relwidth=0.6,relheight=0.6,anchor=N)
Instructions = Label(frame, text= "Select two word SS table .docx files to merge", font="Arial")
Instructions.grid(columnspan=3, column=1, row=1)
ss1_btn = Button(frame, text="Browse SS1", command=lambda:First_browse(),font="Railway", bg="blue", fg="white", height=2, width=11)
ss1_btn.grid(column=1, row=3)
first_text = Text(frame, width=40, height=2,borderwidth=2)
first_text. grid(column=3, row=3,stick=W)
first_text.config(state="disabled")
ss2_btn = Button(frame, text="Browse SS2",command=lambda:Second_browse(), font="Railway", bg="blue", fg="white", height=2, width=11)
ss2_btn.grid(column=1, row=4,columnspan=3,sticky= W)
second_text = Text(frame, width=40, height=2,borderwidth=2)
second_text. grid(column=3, row=4)
second_text.config(state="disabled")
ss3_btn = Button(frame, text="Browse SS3",command=lambda:third_browse(), font="Railway", bg="blue", fg="white", height=2, width=11)
ss3_btn.grid(column=1, row=5)
third_text = Text(frame, width=40, height=2,borderwidth=2)
third_text. grid(column=3, row=5)
third_text.config(state="disabled")
merge_btn = Button(frame, text="Merge", command=lambda: Merge(),font="Railway", bg="blue", fg="white", height=2, width=15)
merge_btn.grid(column=3, row=7)
canvas= tkinter.Canvas(main,width=900, height=400)
canvas.grid(columnspan=6,rowspan=4)
main.mainloop()